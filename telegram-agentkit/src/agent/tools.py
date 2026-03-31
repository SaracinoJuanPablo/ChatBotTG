# src/agent/tools.py
# Herramientas GENÉRICAS para operar sobre archivos.
# No asume nada sobre el contenido — funciona con cualquier archivo
# que el usuario suba a data/knowledge/.
#
# Soporta: .txt, .md, .xlsx, .xlsm, .docx

import logging
from pathlib import Path
from src.config import config
from src.agent.knowledge import knowledge_base

logger = logging.getLogger("agentkit")


# ─────────────────────────────────────────────────────────────────────────────
# SCHEMA DE TOOLS — formato OpenAI/Groq
# ─────────────────────────────────────────────────────────────────────────────

TOOLS_SCHEMA = [
    # ── TXT / MARKDOWN ────────────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "leer_archivo_txt",
            "description": "Lee y devuelve el contenido completo de un archivo .txt o .md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string", "description": "Nombre del archivo, ej: gastos.txt"}
                },
                "required": ["archivo"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "agregar_linea_txt",
            "description": "Agrega una línea al final de un archivo .txt o .md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string"},
                    "texto":   {"type": "string", "description": "Texto de la línea a agregar"}
                },
                "required": ["archivo", "texto"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "eliminar_linea_txt",
            "description": "Elimina todas las líneas que contengan el texto indicado de un archivo .txt o .md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string"},
                    "texto":   {"type": "string", "description": "Texto a buscar y eliminar"}
                },
                "required": ["archivo", "texto"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "reemplazar_contenido_txt",
            "description": "Reemplaza o vacía el contenido completo de un archivo .txt o .md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo":   {"type": "string"},
                    "contenido": {"type": "string", "description": "Nuevo contenido completo. Vacío para vaciar el archivo."}
                },
                "required": ["archivo", "contenido"]
            }
        }
    },
    # ── EXCEL (.xlsx / .xlsm) ─────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "leer_excel",
            "description": "Lee el contenido de un archivo Excel (.xlsx o .xlsm) y lo devuelve como texto con número de fila.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string"},
                    "hoja":    {"type": "string", "description": "Nombre de la hoja (opcional, por defecto la primera)"}
                },
                "required": ["archivo"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "modificar_celda_excel",
            "description": (
                "Modifica el valor de una celda en un Excel. "
                "Busca la fila por un texto y actualiza la columna indicada. "
                "Usar para: cambiar un valor, agregar info a un campo existente, "
                "actualizar un dato, corregir un error en el Excel."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo":     {"type": "string"},
                    "buscar":      {"type": "string", "description": "Texto para identificar la fila (ej: 'Juan Pérez', 'Enero', 'Factura 001')"},
                    "columna":     {"type": "string", "description": "Nombre del encabezado o número de columna (ej: 'Monto', '3')"},
                    "valor_nuevo": {"type": "string", "description": "Nuevo valor para la celda"},
                    "hoja":        {"type": "string", "description": "Nombre de la hoja (opcional)"}
                },
                "required": ["archivo", "buscar", "columna", "valor_nuevo"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "agregar_fila_excel",
            "description": "Agrega una fila nueva al final de un Excel.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo":  {"type": "string"},
                    "valores":  {"type": "array", "items": {"type": "string"}, "description": "Lista de valores para cada columna"},
                    "hoja":     {"type": "string", "description": "Nombre de la hoja (opcional)"}
                },
                "required": ["archivo", "valores"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "eliminar_fila_excel",
            "description": "Elimina filas de un Excel que contengan el texto indicado en cualquier celda.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string"},
                    "texto":   {"type": "string", "description": "Texto a buscar en las filas para eliminar"},
                    "hoja":    {"type": "string", "description": "Nombre de la hoja (opcional)"}
                },
                "required": ["archivo", "texto"]
            }
        }
    },
    # ── WORD (.docx) ──────────────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "leer_word",
            "description": "Lee el contenido de un archivo Word (.docx).",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string"}
                },
                "required": ["archivo"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "agregar_parrafo_word",
            "description": "Agrega un párrafo al final de un archivo Word (.docx).",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo": {"type": "string"},
                    "texto":   {"type": "string"}
                },
                "required": ["archivo", "texto"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "reemplazar_texto_word",
            "description": "Busca un texto en un Word (.docx) y lo reemplaza por otro.",
            "parameters": {
                "type": "object",
                "properties": {
                    "archivo":     {"type": "string"},
                    "buscar":      {"type": "string"},
                    "reemplazar":  {"type": "string"}
                },
                "required": ["archivo", "buscar", "reemplazar"]
            }
        }
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# DISPATCHER
# ─────────────────────────────────────────────────────────────────────────────

def ejecutar_tool(nombre: str, args: dict) -> str:
    dispatch = {
        "leer_archivo_txt":        _leer_archivo_txt,
        "agregar_linea_txt":       _agregar_linea_txt,
        "eliminar_linea_txt":      _eliminar_linea_txt,
        "reemplazar_contenido_txt":_reemplazar_contenido_txt,
        "leer_excel":              _leer_excel,
        "modificar_celda_excel":   _modificar_celda_excel,
        "agregar_fila_excel":      _agregar_fila_excel,
        "eliminar_fila_excel":     _eliminar_fila_excel,
        "leer_word":               _leer_word,
        "agregar_parrafo_word":    _agregar_parrafo_word,
        "reemplazar_texto_word":   _reemplazar_texto_word,
    }
    fn = dispatch.get(nombre)
    if not fn:
        return f"Tool desconocida: '{nombre}'"
    try:
        return fn(**args)
    except TypeError as e:
        return f"Argumentos incorrectos para {nombre}: {e}"
    except Exception as e:
        logger.error(f"Error en tool {nombre}: {e}")
        return f"Error ejecutando {nombre}: {e}"


# ─────────────────────────────────────────────────────────────────────────────
# UTILIDADES INTERNAS
# ─────────────────────────────────────────────────────────────────────────────

def _resolver_ruta(nombre: str) -> Path | None:
    """Resuelve el path de un archivo dentro de knowledge_dir. Bloquea path traversal."""
    for c in ["..", "/", "\\", "\x00"]:
        if c in nombre:
            return None
    ruta = config.knowledge_dir / nombre
    try:
        ruta.resolve().relative_to(config.knowledge_dir.resolve())
    except ValueError:
        return None
    return ruta


def _abrir_excel(ruta: Path):
    import openpyxl
    keep_vba = str(ruta).lower().endswith('.xlsm')
    return openpyxl.load_workbook(ruta, keep_vba=keep_vba)


# ─────────────────────────────────────────────────────────────────────────────
# TXT / MARKDOWN
# ─────────────────────────────────────────────────────────────────────────────

def _leer_archivo_txt(archivo: str) -> str:
    contenido = knowledge_base.ver_archivo(archivo)
    if contenido is None:
        return f"No se encontró '{archivo}'."
    if not contenido.strip():
        return f"'{archivo}' está vacío."
    return contenido


def _agregar_linea_txt(archivo: str, texto: str) -> str:
    ok, msg = knowledge_base.agregar_linea(archivo, texto)
    return msg


def _eliminar_linea_txt(archivo: str, texto: str) -> str:
    ok, msg = knowledge_base.eliminar_lineas(archivo, texto)
    return msg


def _reemplazar_contenido_txt(archivo: str, contenido: str) -> str:
    ok, msg = knowledge_base.reemplazar_contenido(archivo, contenido)
    return msg


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL
# ─────────────────────────────────────────────────────────────────────────────

def _leer_excel(archivo: str, hoja: str = None) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Error: ejecutá 'pip install openpyxl'"

    ruta = _resolver_ruta(archivo)
    if not ruta or not ruta.exists():
        return f"No se encontró '{archivo}'."
    try:
        wb = _abrir_excel(ruta)
        ws = wb[hoja] if hoja and hoja in wb.sheetnames else wb.active
        filas = []
        for i, fila in enumerate(ws.iter_rows(values_only=True), 1):
            if any(c is not None for c in fila):
                filas.append(f"Fila {i}: " + " | ".join(str(c) if c is not None else "" for c in fila))
        if not filas:
            return f"'{archivo}' hoja '{ws.title}' está vacío."
        return f"{archivo} — hoja '{ws.title}':\n" + "\n".join(filas)
    except Exception as e:
        return f"Error leyendo '{archivo}': {e}"


def _modificar_celda_excel(archivo: str, buscar: str, columna: str, valor_nuevo: str, hoja: str = None) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Error: ejecutá 'pip install openpyxl'"

    ruta = _resolver_ruta(archivo)
    if not ruta or not ruta.exists():
        return f"No se encontró '{archivo}'."
    try:
        wb = _abrir_excel(ruta)
        ws = wb[hoja] if hoja and hoja in wb.sheetnames else wb.active

        # Resolver índice de columna — por nombre o por número
        col_idx = None
        if columna.isdigit():
            col_idx = int(columna)
        else:
            for cell in next(ws.iter_rows(max_row=1)):
                if cell.value and columna.lower() in str(cell.value).lower():
                    col_idx = cell.column
                    break

        if col_idx is None:
            # Si no encontró por encabezado, mostrar encabezados disponibles
            headers = [str(c.value) for c in next(ws.iter_rows(max_row=1)) if c.value]
            return f"No se encontró columna '{columna}'. Columnas disponibles: {', '.join(headers)}"

        # Buscar y modificar filas
        buscar_lower = buscar.lower()
        modificadas = 0
        for fila in ws.iter_rows():
            for celda in fila:
                if celda.value and buscar_lower in str(celda.value).lower():
                    ws.cell(row=celda.row, column=col_idx, value=valor_nuevo)
                    modificadas += 1
                    break

        if modificadas == 0:
            return f"No se encontró '{buscar}' en '{archivo}'."

        wb.save(ruta)
        return f"'{archivo}': columna '{columna}' actualizada a '{valor_nuevo}' en {modificadas} fila(s) que contenían '{buscar}'."
    except Exception as e:
        return f"Error modificando '{archivo}': {e}"


def _agregar_fila_excel(archivo: str, valores: list, hoja: str = None) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Error: ejecutá 'pip install openpyxl'"

    ruta = _resolver_ruta(archivo)
    if not ruta:
        return f"Archivo inválido: '{archivo}'."
    try:
        if ruta.exists():
            wb = _abrir_excel(ruta)
            ws = wb[hoja] if hoja and hoja in wb.sheetnames else wb.active
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
        ws.append(valores)
        wb.save(ruta)
        return f"Fila agregada en '{archivo}': {' | '.join(str(v) for v in valores)}"
    except Exception as e:
        return f"Error en '{archivo}': {e}"


def _eliminar_fila_excel(archivo: str, texto: str, hoja: str = None) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Error: ejecutá 'pip install openpyxl'"

    ruta = _resolver_ruta(archivo)
    if not ruta or not ruta.exists():
        return f"No se encontró '{archivo}'."
    try:
        wb = _abrir_excel(ruta)
        ws = wb[hoja] if hoja and hoja in wb.sheetnames else wb.active
        texto_lower = texto.lower()
        a_eliminar = []
        for fila in ws.iter_rows():
            for celda in fila:
                if celda.value and texto_lower in str(celda.value).lower():
                    a_eliminar.append(celda.row)
                    break
        if not a_eliminar:
            return f"No se encontró '{texto}' en '{archivo}'."
        for num in sorted(set(a_eliminar), reverse=True):
            ws.delete_rows(num)
        wb.save(ruta)
        return f"Eliminadas {len(a_eliminar)} fila(s) con '{texto}' de '{archivo}'."
    except Exception as e:
        return f"Error en '{archivo}': {e}"


# ─────────────────────────────────────────────────────────────────────────────
# WORD
# ─────────────────────────────────────────────────────────────────────────────

def _leer_word(archivo: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: ejecutá 'pip install python-docx'"
    ruta = _resolver_ruta(archivo)
    if not ruta or not ruta.exists():
        return f"No se encontró '{archivo}'."
    try:
        doc = Document(str(ruta))
        parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        return f"{archivo}:\n\n" + "\n".join(parrafos) if parrafos else f"'{archivo}' está vacío."
    except Exception as e:
        return f"Error leyendo '{archivo}': {e}"


def _agregar_parrafo_word(archivo: str, texto: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: ejecutá 'pip install python-docx'"
    ruta = _resolver_ruta(archivo)
    if not ruta:
        return f"Archivo inválido: '{archivo}'."
    try:
        doc = Document(str(ruta)) if ruta.exists() else Document()
        doc.add_paragraph(texto)
        doc.save(str(ruta))
        return f"Párrafo agregado en '{archivo}'."
    except Exception as e:
        return f"Error en '{archivo}': {e}"


def _reemplazar_texto_word(archivo: str, buscar: str, reemplazar: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: ejecutá 'pip install python-docx'"
    ruta = _resolver_ruta(archivo)
    if not ruta or not ruta.exists():
        return f"No se encontró '{archivo}'."
    try:
        doc = Document(str(ruta))
        encontrados = 0
        for p in doc.paragraphs:
            if buscar in p.text:
                for run in p.runs:
                    if buscar in run.text:
                        run.text = run.text.replace(buscar, reemplazar)
                        encontrados += 1
        if encontrados == 0:
            return f"No se encontró '{buscar}' en '{archivo}'."
        doc.save(str(ruta))
        return f"Reemplazadas {encontrados} ocurrencia(s) de '{buscar}' → '{reemplazar}' en '{archivo}'."
    except Exception as e:
        return f"Error en '{archivo}': {e}"